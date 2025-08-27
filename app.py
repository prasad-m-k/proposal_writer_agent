# app.py

import os
import pandas as pd
import numpy as np
import re
import time
from flask import Flask, render_template, request, url_for, send_from_directory, abort
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from convert import MarkdownToDocxConverter # Your updated converter
from datetime import date
# from flask_wtf.csrf import CSRFProtect # REMOVED CSRF
from werkzeug.utils import secure_filename


app = Flask(__name__)
env = {}
# csrf = CSRFProtect() # REMOVED CSRF

def _initialize():
    """Loads environment variables and configures the Flask app."""
    global env
    load_dotenv()
    
    # --- Security Configuration ---
    # SECRET_KEY is not strictly needed for CSRF anymore, but kept for other Flask session management if any
    app.config['SECRET_KEY'] = os.getenv("SECRET_KEY", "a_super_secret_key_for_dev_if_not_set")


    # --- App Configuration ---
    UPLOAD_FOLDER = 'uploads'
    app.config['UPLOAD_FOLDER'] = os.path.abspath(UPLOAD_FOLDER)
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
        
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not set. Please set it in your .env file.")
    genai.configure(api_key=GEMINI_API_KEY)
    
    env['DEPL'] = os.getenv("DEPL")


# Security Headers Configuration
@app.after_request
def apply_security_headers(response):
    """Applies security-enhancing HTTP headers to every response."""
    # --- CORRECTED CONTENT SECURITY POLICY ---
    # This policy now correctly allows styles and fonts from Google, restoring the UI.
    # Added 'unsafe-inline' to style-src to allow inline <style> tags and style attributes.
    # Added 'data:' to font-src for broader font support, though https://fonts.gstatic.com covers Google Fonts.
    csp = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline'; " # 'unsafe-inline' is often needed for scripts in templates
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; " # ADDED 'unsafe-inline'
        "font-src 'self' https://fonts.gstatic.com data:; " # ADDED 'data:'
    )
    response.headers['Content-Security-Policy'] = csp
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-Content-Type-Options'] = 'nosniff'
    if env.get('DEPL') == "PROD":
        response.headers['Strict-TransportSecurity'] = 'max-age=31536000; includeSubDomains'
    return response

def create_document_header(document):
    """
    Adds a pre-defined, crisp header with a smaller font to a document object.
    """
    LOGO_PATH = 'static/assets/msg_logo.png'
    ADDRESS_LINES = [
        ("Musical Instruments N Kids Hands", True, 8),
        ("Music Science Group", True, 10),
        ("2150 Capitol Avenue Sacramento, CA 95816", False, 8),
        ("Ph. (916) 670-9950", False, 8)
    ]

    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""

    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.5)
    header_table.columns[1].width = Inches(5.0)

    logo_cell = header_table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(LOGO_PATH, height=Inches(0.75))

    address_cell = header_table.cell(0, 1)
    address_paragraph = address_cell.paragraphs[0]
    address_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for text, is_bold, font_size in ADDRESS_LINES:
        run = address_paragraph.add_run(text + '\n')
        run.bold = is_bold
        font = run.font
        font.size = Pt(font_size)

def normalize_empty_lines(text):
    """
    Reduces any sequence of three or more newlines to exactly two newlines.
    """
    return re.sub(r'\n{2,}', r'\n', text)

def generate_proposal_from_row(district="N/A", cost_proposal="N/A", num_weeks="N/A", days_per_week="N/A", selected_schools=[]):
    """
    Generates a proposal with a custom header, saves it as a .docx file,
    and returns the AI text and the filename.
    """
    try:
        with open('uploads/district.csv', 'r') as f:
            all_districts_info = f.read()

        with open('uploads/about_msg.txt', 'r') as f:
            about_msg = f.read()

        with open('uploads/proposal.txt', 'r') as file:
            proposal_context = file.read()

        model = genai.GenerativeModel('gemini-1.5-flash')
        company = "Music Science Group"
        client = f"{district} School District"
        project = "Educational Learning Outreach Program (ELOP)"
        services = "Music Integration, S.T.E.A.M. Education, Wellness Programs, Student Engagement Activities"
        benefits = "Enhanced Student Engagement, Improved Academic Performance, Increased Wellness, Community Involvement"
        cta = "We look forward to the opportunity to collaborate and make a meaningful impact together."
        today = date.today()
        
        school_locations = ", ".join(selected_schools) if selected_schools else "Selected school sites"

        prompt = f"""
        You are a professional proposal writer for Music Science Group (MSG) / Musical Instruments N Kids Hands (M.I.N.K.H.). Your goal is to generate a comprehensive, client-centric project proposal for the {district} School District.

        **Here is the essential information you need for the proposal:**
        *   **Company Name:** Musical Instruments N Kids Hands (M.I.N.K.H.) / Music Science Group (MSG)
        *   **Client Name:** {district} School District
        *   **Project Name:** Educational Learning Outreach Program (ELOP)
        *   **Services Offered:** Music Integration, S.T.E.A.M. Education, Wellness Programs, Student Engagement Activities
        *   **Key Benefits:** Enhanced Student Engagement, Improved Academic Performance, Increased Wellness, Community Involvement
        *   **Call to Action Phrase:** "We look forward to the opportunity to collaborate and make a meaningful impact together."
        *   **Today's Date:** {today}
        *   **Program Schedule:** {days_per_week} day(s) per week
        *   **Program Duration:** {num_weeks} weeks
        *   **Program Locations:** {school_locations} (e.g., "Selected school sites" or a comma-separated list)
        *   **Total Program Fee:** ${float(cost_proposal):,.2f}
        *   **Contact Person:** A.P. Moore, Programs Coordinator
        *   **Contact Email:** aphilanda@musicsciencegroup.com
        *   **Grades Served:** TK-8th grade
        *   **Participant Capacity:** Maximum 40 students per class, with a total enrollment capacity of 600 students.

        **Contextual Information (use this to inform the content of each section):**
        *   **About MSG/MINKH:** "{about_msg}"
        *   **General Proposal Details/Program Overview:** "{proposal_context}"
        *   **District-Specific Data:** (Assume `all_districts_info` is used internally by the model to tailor responses, if needed. No need to explicitly include its content in the prompt, just acknowledge its availability).

        **Proposal Structure and Content Requirements:**

        The proposal must be structured with clear, bold markdown headings (e.g., **1. Executive Summary**). Maintain appropriate spacing and alignment for a formal business document; avoid excessive empty lines between paragraphs or sections.

        1.  **Header Information:**
            *   **Date:** {today}
            *   **Submitted to:** {district} School District
            *   **Submitted by:** Musical Instruments N Kids Hands
            *   **Contact:**
                *   A.P. Moore, Programs Coordinator
                *   Email: aphilanda@musicsciencegroup.com

        2.  **Introduction:** Provide a welcoming overview, briefly stating the purpose of the proposal and the partnership with {client}.

        3.  **Executive Summary:** A concise and compelling overview.
            *   Highlight {client}'s potential challenges (derived from `problem_statement`).
            *   Present MSG/MINKH's proposed solution ({project}).
            *   Outline the key, measurable benefits for {client} ({benefits}), referencing the innovative program offerings.

        4.  **Problem Statement:** Clearly and empathetically articulate the specific challenges and needs of {client} that MSG/MINKH aims to address, drawing from the `proposal_context` and `about_msg`.

        5.  **Proposed Solution - {project}:** Detail the specific services and approach MSG/MINKH will deploy.
            *   Explain how the comprehensive services ({services}) directly address the problems identified in the Problem Statement.
            *   Emphasize the unique value of MSG/MINKH's approach.
            *   **Program Offerings:** Describe how MSG's enriched curriculum transforms classrooms into dynamic centers where music and technology converge. Specifically highlight:
                *   Digital music stations that turn creativity into sound
                *   3D printing of instruments that bring design and innovation to life
                *   Artificial Intelligence tools that make music creation interactive and exciting
                *   Robotics, coding, and programming that build skills for tomorrow’s careers
                *   Entry-level engineering projects that teach problem-solving and innovation
                *   Digital arts and media creation to inspire self-expression and storytelling
                *   Introductory music instrument instruction that sparks joy and confidence

        6.  **Benefits and Value Proposition:** Elaborate on the tangible and intangible benefits {client} will gain. Specifically emphasize: {benefits}. Quantify or illustrate benefits where possible, linking them to the innovative program offerings.

        7.  **Teamwork and Collaboration:** Describe how MSG/MINKH will partner with {client} staff, teachers, and the community to ensure program success and integration.

        8.  **Music for Optimal Wellness:** Dedicate a section to how the program specifically promotes student wellness through music, drawing from `about_msg` and `proposal_context`, and how the interactive offerings contribute to this.

        9.  **Program Details:**
            *   **Schedule:** {days_per_week} day(s) per week.
            *   **Duration:** {num_weeks} weeks.
            *   **Dates:** Provide an example start date (e.g., "August 13, {year}") and calculate an estimated end date based on the duration.
            *   **Locations:** {school_locations}
            *   **Grades Served:** The program is designed for students in grades TK-8th.
            *   **Participant Capacity:** Each class can accommodate a maximum of 40 students, with a total program enrollment capacity of 600 students across all sites. (Referencing a "lesson plan" should be handled as an external document, not something the AI generates in detail here.)

        10. **Fee:** Clearly state: "The total fee for the program is ${float(cost_proposal):,.2f}."

        11. **Payment Schedule:** Propose a standard payment schedule (e.g., "50% upfront, 50% upon completion," or "monthly installments"). *Self-correction: If no specific schedule is provided, create a reasonable placeholder.*

        12. **Call to Action:** A clear and persuasive statement encouraging next steps. Include the exact phrase: "{cta}".

        13. **Conclusion:** Summarize the proposal's core message, reiterate commitment to {client}'s success, and express enthusiasm for a partnership.

        **Formatting Guidelines:**
        *   Use professional, client-centric language throughout.
        *   Utilize markdown for headings, bolding, and italics to enhance readability and emphasis.
        *   Ensure clarity, conciseness, and a logical flow.
        *   Avoid conversational filler; go straight to the proposal content after the initial header information.
        """
        
        print("Generating proposal using Gemini model...")
        response = model.generate_content(prompt)
        text = normalize_empty_lines(response.text)
        
        document = Document()
        create_document_header(document)
        converter = MarkdownToDocxConverter(document=document)
        converter.convert(text)

        safe_district_name = re.sub(r'[^a-zA-Z0-9_]', '', district).replace(" ", "_")
        timestamp = int(time.time())
        filename = f"Proposal_{safe_district_name}_{timestamp}.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        document.save(output_path)
        print(f"Successfully created '{filename}' with custom header.")

        return text, filename
    except Exception as e:
        error_text = f"An error occurred while generating the proposal: {e}"
        print(error_text)
        return error_text, None

def get_school_data():
    df = pd.read_csv("data.csv")
    df.columns = df.columns.str.strip()
    return df

@app.route('/')
def index():
    df = get_school_data()
    districts = df['District Name'].unique().tolist()
    all_data_json = df.to_json(orient='records')
    return render_template('index.html', districts=districts, all_data=all_data_json)

@app.route('/proposal', methods=['POST'])
def generate_proposal():
    district_name = request.form.get('district')
    cost_proposal = request.form.get('cost_proposal')
    num_weeks = request.form.get('num_weeks')
    days_per_week = request.form.get('days_per_week')
    selected_schools = request.form.getlist('schoolname')

    proposal_text, filename = generate_proposal_from_row(
        district=district_name, 
        cost_proposal=cost_proposal, 
        num_weeks=num_weeks, 
        days_per_week=days_per_week, 
        selected_schools=selected_schools
    )

    proposal_data = {
        'district': district_name,
        'cost_proposal': cost_proposal,
        'school_name': ", ".join(selected_schools),
        'proposal_text': proposal_text,
        'filename': filename
    }
    return render_template('proposal.html', data=proposal_data)

@app.route('/download/<path:filename>')
def download_file(filename):
    safe_filename = secure_filename(filename)
    upload_dir = app.config['UPLOAD_FOLDER']
    file_path = os.path.join(upload_dir, safe_filename)
    
    if not os.path.normpath(file_path).startswith(upload_dir) or not os.path.isfile(file_path):
        abort(404)

    return send_from_directory(upload_dir, safe_filename, as_attachment=True)

if __name__ == '__main__':
    # Initialize the application
    _initialize()

    if env.get("DEPL") == "PROD":
        app.run(host='0.0.0.0', port=443, debug=False, ssl_context=(
            '/etc/letsencrypt/live/qaproposalmusicsciencegroup.com/fullchain.pem',
            '/etc/letsencrypt/live/qaproposalmusicsciencegroup.com/privkey.pem'
        ))
    else:
        app.run(port=5001, debug=True)
