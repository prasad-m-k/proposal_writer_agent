"""
Word Document Formatter Module

This module provides functions to format Word documents by:
1. Converting **text** to bold formatting
2. Adding spaces after colons (text:word -> text: word)
"""

from docx import Document
import re

def format_word_document(input_file, output_file):
    """
    Scan a Word document and:
    1. Make text between ** bold
    2. Add space after colons where missing
    
    Args:
        input_file (str): Path to input Word document
        output_file (str): Path to save formatted document
    """
    # Load the document
    document = Document(input_file)
    
    # Process all paragraphs
    for paragraph in document.paragraphs:
        format_paragraph(paragraph)
    
    # Process tables if they exist
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    format_paragraph(paragraph)
    
    # Save the modified document
    document.save(output_file)
    print(f"Document formatted and saved as: {output_file}")

def format_paragraph(paragraph):
    """
    Format a single paragraph by processing its runs
    
    Args:
        paragraph: python-docx paragraph object
    """
    if not paragraph.text.strip():
        return
    
    # Get the full text of the paragraph
    full_text = paragraph.text
    
    # Check if formatting is needed
    needs_bold_formatting = '**' in full_text
    needs_colon_spacing = re.search(r'\w:[^\s]', full_text)
    
    if not (needs_bold_formatting or needs_colon_spacing):
        return
    
    # Clear existing runs
    paragraph.clear()
    
    # Apply colon spacing first
    if needs_colon_spacing:
        full_text = re.sub(r'(\w):([^\s])', r'\1: \2', full_text)
    
    # Process bold formatting
    if needs_bold_formatting:
        process_bold_formatting(paragraph, full_text)
    else:
        # Just add the text with colon spacing
        paragraph.add_run(full_text)

def process_bold_formatting(paragraph, text):
    """
    Process text to identify bold sections and add them as runs
    
    Args:
        paragraph: python-docx paragraph object
        text (str): Text to process for bold formatting
    """
    # Split text by ** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            # This is bold text (remove the ** markers)
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif part:  # Not empty
            # Regular text
            paragraph.add_run(part)

def format_word_document_advanced(input_file, output_file):
    """
    Advanced version that preserves existing formatting while adding new formatting
    
    Args:
        input_file (str): Path to input Word document
        output_file (str): Path to save formatted document
    """
    document = Document(input_file)
    
    # Process all paragraphs
    for paragraph in document.paragraphs:
        format_paragraph_advanced(paragraph)
    
    # Process tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    format_paragraph_advanced(paragraph)
    
    document.save(output_file)
    print(f"Document formatted (advanced) and saved as: {output_file}")

def format_paragraph_advanced(paragraph):
    """
    Advanced formatting that preserves existing run properties
    
    Args:
        paragraph: python-docx paragraph object
    """
    if not paragraph.text.strip():
        return
    
    full_text = paragraph.text
    needs_formatting = '**' in full_text or re.search(r'\w:[^\s]', full_text)
    
    if not needs_formatting:
        return
    
    # Store original paragraph formatting
    original_style = paragraph.style
    original_alignment = paragraph.alignment
    
    # Get the first run's font properties to preserve them
    font_properties = {}
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_properties = {
            'name': first_run.font.name,
            'size': first_run.font.size,
            'color': first_run.font.color.rgb if first_run.font.color.rgb else None,
            'italic': first_run.italic,
            'underline': first_run.underline
        }
    
    # Clear and rebuild
    paragraph.clear()
    paragraph.style = original_style
    paragraph.alignment = original_alignment
    
    # Apply colon spacing
    if re.search(r'\w:[^\s]', full_text):
        full_text = re.sub(r'(\w):([^\s])', r'\1: \2', full_text)
    
    # Process bold formatting
    parts = re.split(r'(\*\*.*?\*\*)', full_text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            # Bold text
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
            # Apply preserved font properties
            apply_font_properties(run, font_properties)
        elif part:
            # Regular text
            run = paragraph.add_run(part)
            apply_font_properties(run, font_properties)

def apply_font_properties(run, font_properties):
    """
    Apply font properties to a run
    
    Args:
        run: python-docx run object
        font_properties (dict): Dictionary containing font properties
    """
    if font_properties.get('name'):
        run.font.name = font_properties['name']
    if font_properties.get('size'):
        run.font.size = font_properties['size']
    if font_properties.get('color'):
        run.font.color.rgb = font_properties['color']
    if font_properties.get('italic'):
        run.italic = font_properties['italic']
    if font_properties.get('underline'):
        run.underline = font_properties['underline']

def color_headings(docx_path):
    """Opens a DOCX file and changes the font color of all heading styles to blue."""
    #return
    try:
        document = Document(docx_path)
        styles = document.styles

        # Find the built-in heading styles and set their font color
        heading_styles = {
            'Heading 1': styles['Heading 1'],
            'Heading 2': styles['Heading 2'],
            'Heading 3': styles['Heading 3'],
            # Add more as needed
        }

        blue = RGBColor(0x00, 0x00, 0xFF) # RGB for blue

        for style in heading_styles.values():
            if style:
                style.font.color.rgb = blue
            
        # Iterate through the document paragraphs and apply the style to any headings
        for paragraph in document.paragraphs:
            if paragraph.style.name in heading_styles:
                paragraph.style.font.color.rgb = blue

        document.save("blue_headings_output.docx")
        print("Successfully updated heading colors and saved to 'blue_headings_output.docx'.")

    except Exception as e:
        print(f"An error occurred while modifying the document: {e}")

def test_formatting():
    """
    Test function to demonstrate the formatting logic
    """
    test_cases = [
        "This is **bold text** in a sentence.",
        "Multiple **bold** words and **more bold** text.",
        "Name:John should become Name: John",
        "Time:15:30 and Date:2024-01-01",
        "Mixed **bold** and Name:John formatting.",
        "**Bold at start** and end is **bold**"
    ]
    
    print("Testing formatting logic:")
    for test in test_cases:
        print(f"Original: {test}")
        # Simulate the colon spacing
        formatted = re.sub(r'(\w):([^\s])', r'\1: \2', test)
        # Show bold markers (in actual document, these become bold formatting)
        print(f"Formatted: {formatted}")
        print("-" * 50)
