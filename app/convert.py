# convert.py

import re
from docx import Document
from docx.shared import Inches, Pt

class MarkdownToDocxConverter:
    """
    A class to convert a Markdown formatted string into a .docx file.
    Can be initialized with an existing Document object to append content.
    """

    def __init__(self, document=None):
        """
        Initializes the converter.
        Args:
            document (Document, optional): An existing python-docx Document object. 
                                           If None, a new one is created.
        """
        if document:
            self.doc = document
        else:
            self.doc = Document()

    def convert(self, markdown_text):
        """
        Main conversion method. Adds markdown text to the document object.
        Note: This method no longer saves the file.

        Args:
            markdown_text (str): The string containing Markdown text.
        """
        self.inline_pattern = re.compile(
            r'(\*{1,3}|_{1,3}|~~|`)'
            r'(.+?)'
            r'\1'
        )
        lines = markdown_text.strip().split('\n')

        # Track table state
        table_lines = []
        in_table = False

        for line in lines:
            # Check if this line is part of a table
            if self._is_table_line(line):
                table_lines.append(line)
                in_table = True
                continue
            elif in_table:
                # We've reached the end of a table
                self._add_table(table_lines)
                table_lines = []
                in_table = False

            if line.startswith('#'):
                self._add_heading(line)
            elif line.strip() in ['---', '***', '___']:
                self._add_horizontal_rule()
            elif line.startswith('>'):
                self._add_blockquote(line)
            elif re.match(r'^\s*(\*|-|\+)\s', line) or re.match(r'^\s*\d+\.\s', line):
                self._add_list_item(line)
            elif line.strip() == '':
                # Skip adding paragraphs for empty lines to reduce spacing
                pass
            else:
                self._add_formatted_paragraph(line)

        # Handle table at end of document
        if in_table and table_lines:
            self._add_table(table_lines)
        # The document is NOT saved here. The calling function is responsible for saving.

    # ... (the rest of the _add_... methods remain exactly the same) ...

    def _add_heading(self, line):
        """Adds a heading based on the number of '#' characters."""
        level = 0
        while line[level] == '#':
            level += 1
        
        heading_text = line[level:].strip()
        self.doc.add_heading(heading_text, level=min(level, 6))

    def _add_horizontal_rule(self):
        """Adds a horizontal rule to the document."""
        # Add a simple line break instead of underscores
        p = self.doc.add_paragraph()
        # Add some spacing by creating an empty paragraph
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def _add_blockquote(self, line):
        """Adds a blockquote, removing the '>' marker."""
        quote_text = line.lstrip('> ').strip()
        self.doc.add_paragraph(quote_text, style='Quote')

    def _add_list_item(self, line):
        """Adds a list item, handling indentation and type (bullet/number)."""

        indent_level = len(line) - len(line.lstrip())
        indent_spaces = indent_level // 2

        # Determine base style
        if re.match(r'^\s*\d+\.\s', line):
            base_style = 'List Number'
        else:
            base_style = 'List Bullet'

        # Limit indentation levels to available styles (typically 1-3)
        if indent_spaces > 0:
            # Cap the indentation level to avoid non-existent styles
            level = min(indent_spaces + 1, 3)
            if level == 2:
                style = f"{base_style} 2"
            elif level == 3:
                style = f"{base_style} 3"
            else:
                style = base_style
        else:
            style = base_style
        # Fixed regex: properly handle numbered lists with single space
        text = re.sub(r'^\s*(\*|-|\+)\s*|^\s*\d+\.\s', '', line).strip()


        # Try to add paragraph with the requested style, fallback to basic list style if it doesn't exist
        try:
            self.doc.add_paragraph(text, style=style)
        except KeyError:
            # If the specific style doesn't exist, fall back to the base style
            fallback_style = base_style if 'base_style' in locals() else 'List Bullet'
            try:
                self.doc.add_paragraph(text, style=fallback_style)
            except KeyError:
                # If even the base style doesn't exist, create as normal paragraph
                paragraph = self.doc.add_paragraph(text)
                # Add bullet formatting manually if needed
                if not re.match(r'^\s*\d+\.\s', line):
                    paragraph.style = 'Normal'

    def _add_formatted_paragraph(self, line):
        """Adds a paragraph with complex inline formatting."""
        p = self.doc.add_paragraph()

        # Reduce spacing for form fields (lines with **Field:** pattern)
        if re.match(r'^\*\*[^:]+:\*\*', line):
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        last_end = 0
        for match in self.inline_pattern.finditer(line):
            start, end = match.span()
            marker = match.group(1)
            text = match.group(2)
            if start > last_end:
                p.add_run(line[last_end:start])
            run = p.add_run(text)
            if marker in ['**', '__', '***', '___']:
                run.bold = True
            if marker in ['*', '_', '***', '___']:
                run.italic = True
            if marker == '`':
                font = run.font
                font.name = 'Courier New'
            if marker == '~~':
                run.font.strike = True
            last_end = end
        if last_end < len(line):
            p.add_run(line[last_end:])

    def _is_table_line(self, line):
        """Check if a line is part of a markdown table"""
        # A table line contains pipe characters (|) and isn't empty
        return '|' in line.strip() and line.strip() != '' and not line.strip().startswith('#')

    def _add_table(self, table_lines):
        """Add a table to the document from markdown table lines"""
        if not table_lines:
            return

        # Filter out separator lines (lines with only |, -, :, and spaces)
        data_lines = []
        for line in table_lines:
            if not re.match(r'^\s*\|?[\s\-:|\|]*\|?\s*$', line):
                data_lines.append(line)

        if not data_lines:
            return

        # Parse table data
        table_data = []
        for line in data_lines:
            # Split by | and clean up cells
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty cells at start/end that come from leading/trailing |
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            if cells:  # Only add non-empty rows
                table_data.append(cells)

        if not table_data:
            return

        # Determine table dimensions
        max_cols = max(len(row) for row in table_data)
        rows = len(table_data)

        # Create the table
        table = self.doc.add_table(rows=rows, cols=max_cols)
        table.style = 'Table Grid'  # Use a standard table style

        # Fill the table
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < max_cols:
                    cell = table.cell(row_idx, col_idx)
                    # Handle bold formatting in cells
                    self._add_formatted_text_to_cell(cell, cell_data)

        # Add some spacing after the table
        self.doc.add_paragraph()

    def _add_formatted_text_to_cell(self, cell, text):
        """Add formatted text to a table cell"""
        # Clear existing content
        cell.text = ''
        paragraph = cell.paragraphs[0]

        # Process inline formatting
        last_end = 0
        for match in self.inline_pattern.finditer(text):
            start, end = match.span()
            marker = match.group(1)
            content = match.group(2)

            # Add text before the match
            if start > last_end:
                paragraph.add_run(text[last_end:start])

            # Add formatted text
            run = paragraph.add_run(content)
            if marker in ['**', '__', '***', '___']:
                run.bold = True
            if marker in ['*', '_', '***', '___']:
                run.italic = True

            last_end = end

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])