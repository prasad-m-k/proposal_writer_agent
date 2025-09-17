# app.py

import os
import pandas as pd
import numpy as np
import re
import time
from flask import Flask, render_template, request, url_for, send_from_directory, abort
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from convert import MarkdownToDocxConverter # Your updated converter
from datetime import date
# from flask_wtf.csrf import CSRFProtect # REMOVED CSRF
from werkzeug.utils import secure_filename
import signal # Added for stopServer
from flask import jsonify # Added for stopServer
from pathlib import Path


app = Flask(__name__)
env = {}
# csrf = CSRFProtect() # REMOVED CSRF

## Define global paths for clarity
INPUT_FILES_FOLDER_NAME= '/app/input_data' # Renamed from 'uploads' for input files
DOWNLOAD_FOLDER_NAME = '/app/generated_proposals' # New folder for output files



def _initialize():
    """Loads environment variables and configures the Flask app."""
    print("Loads environment variables and configures the Flask app")
    global env
    load_dotenv()
    
    # --- Security Configuration ---
    # SECRET_KEY is not strictly needed for CSRF anymore, but kept for other Flask session management if any
    app.config['SECRET_KEY'] = os.getenv("SECRET_KEY", "a_super_secret_key_for_dev_if_not_set")


    #--- App Configuration ---
    # Configure input files folder
    app.config['INPUT_FILES_FOLDER'] = os.path.abspath(INPUT_FILES_FOLDER_NAME)
    if not os.path.exists(app.config['INPUT_FILES_FOLDER']):
        os.makedirs(app.config['INPUT_FILES_FOLDER'])
        
    # Configure download folder for generated documents
    app.config['DOWNLOAD_FOLDER'] = os.path.abspath(DOWNLOAD_FOLDER_NAME)
    if not os.path.exists(app.config['DOWNLOAD_FOLDER']):
        os.makedirs(app.config['DOWNLOAD_FOLDER'])

        '''
    APP_ROOT = Path(__file__).resolve().parent  # directory containing app.py

    # Env can be absolute or relative; relative resolves against APP_ROOT
    _input_env = os.getenv("INPUT_FILES_FOLDER", "input_data")
    INPUT_FILES_FOLDER = Path(_input_env)
    if not INPUT_FILES_FOLDER.is_absolute():
        INPUT_FILES_FOLDER = (APP_ROOT / INPUT_FILES_FOLDER).resolve()

    # Ensure it exists
    INPUT_FILES_FOLDER.mkdir(parents=True, exist_ok=True)
    app.config['INPUT_FILES_FOLDER'] = os.path.abspath(INPUT_FILES_FOLDER)
    print("-------------------------------- inside init -------------------------------")
    print(app.config['INPUT_FILES_FOLDER'], INPUT_FILES_FOLDER)
    print("-------------------------------- inside init -------------------------------")   

    _input_env = os.getenv("DOWNLOAD_FOLDER", "generated_proposals")
    DOWNLOAD_FOLDER = Path(_input_env)
    if not DOWNLOAD_FOLDER.is_absolute():
        DOWNLOAD_FOLDER = (APP_ROOT / DOWNLOAD_FOLDER).resolve()

    # Ensure it exists
    DOWNLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
    app.config['DOWNLOAD_FOLDER'] = os.path.abspath(DOWNLOAD_FOLDER)
    print("-------------------------------- inside init -------------------------------")
    print(app.config['DOWNLOAD_FOLDER'], DOWNLOAD_FOLDER)
    print("-------------------------------- inside init -------------------------------")   
    '''
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not set. Please set it in your .env file.")
    genai.configure(api_key=GEMINI_API_KEY)
    
    env['DEPL'] = os.getenv("DEPL")


# Security Headers Configuration
@app.after_request
def apply_security_headers(response):
    """Applies security-enhancing HTTP headers to every response."""
    # --- CORRECTED CONTENT SECURITY POLICY ---
    # This policy now correctly allows styles and fonts from Google, restoring the UI.
    # Added 'unsafe-inline' to style-src to allow inline <style> tags and style attributes.
    # Added 'data:' to font-src for broader font support, though https://fonts.gstatic.com covers Google Fonts.
    csp = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline'; " # 'unsafe-inline' is often needed for scripts in templates
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; " # ADDED 'unsafe-inline'
        "font-src 'self' https://fonts.gstatic.com data:; " # ADDED 'data:'
    )
    response.headers['Content-Security-Policy'] = csp
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-Content-Type-Options'] = 'nosniff'
    if env.get('DEPL') == "PROD":
        response.headers['Strict-TransportSecurity'] = 'max-age=31536000; includeSubDomains'
    return response

def create_document_header(document):
    """
    Adds a pre-defined, crisp header with a smaller font to a document object.
    """
    LOGO_PATH = 'static/assets/msg_logo.png' # Assuming logo path remains the same
    ADDRESS_LINES = [
        ("Musical Instruments N Kids Hands", True, 8),
        ("Music Science & Technology Group", True, 10), # Updated here
        ("2150 Capitol Avenue Sacramento, CA 95816", False, 8),
        ("Ph. (916) 670-9950", False, 8)
    ]

    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""

    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.5)
    header_table.columns[1].width = Inches(5.0)

    logo_cell = header_table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(LOGO_PATH, height=Inches(0.75))

    address_cell = header_table.cell(0, 1)
    address_paragraph = address_cell.paragraphs[0]
    address_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for text, is_bold, font_size in ADDRESS_LINES:
        run = address_paragraph.add_run(text + '\n')
        run.bold = is_bold
        font = run.font
        font.size = Pt(font_size)

def normalize_empty_lines(text):
    """
    Reduces any sequence of three or more newlines to exactly two newlines.
    """
    return re.sub(r'\n{2,}', r'\n', text)

def manage_file_rotation(district_name, download_folder, max_files=5):
    """
    Manages the rotation of generated proposal files for a given district,
    keeping only the 'max_files' most recent ones.
    """
    safe_district_name = re.sub(r'[^a-zA-Z0-9_]', '', district_name).replace(" ", "_")
    
    # Pattern to match files for this specific district
    file_pattern = rf"Proposal_{re.escape(safe_district_name)}_(\d+)\.docx"
    
    district_files = []
    for filename in os.listdir(download_folder):
        match = re.match(file_pattern, filename)
        if match:
            timestamp = int(match.group(1))
            district_files.append((timestamp, os.path.join(download_folder, filename)))
            
    # Sort files by timestamp, oldest first
    district_files.sort(key=lambda x: x[0])

    # If there are more than max_files, delete the oldest ones
    if len(district_files) > max_files:
        files_to_delete = district_files[:len(district_files) - max_files]
        for _, filepath in files_to_delete:
            try:
                os.remove(filepath)
                print(f"Deleted old proposal file: {filepath}")
            except OSError as e:
                print(f"Error deleting file {filepath}: {e}")

def generate_proposal_from_row(district="N/A", cost_proposal="N/A", num_weeks="N/A", days_per_week="N/A", selected_schools=[]):
    """
    Generates a proposal with a custom header, saves it as a .docx file,
    and returns the AI text and the filename.
    """
    try:
        # Load common contextual files from the INPUT_FILES_FOLDER
        with open(os.path.join(app.config['INPUT_FILES_FOLDER'], 'district.csv'), 'r') as f:
            all_districts_info = f.read()

        with open(os.path.join(app.config['INPUT_FILES_FOLDER'], 'about_msg.txt'), 'r') as f:
            about_msg = f.read()

        with open(os.path.join(app.config['INPUT_FILES_FOLDER'], 'proposal.txt'), 'r') as file:
            proposal_context = file.read()

        # --- NEW: Load Natomas-specific RFP requirements if applicable ---
        natomas_rfp_requirements_context = ""
        natomas_rfp_instruction_formatted = ""
        if district == "Natomas Unified School District":
            try:
                # Assuming this file is in the root or a 'data' folder
                with open('natomas_school_district_rfp_requirements.txt', 'r') as f:
                    natomas_rfp_requirements_context = f.read()
                    natomas_rfp_instruction_formatted = (
                        "\nWhen generating the proposal for Natomas Unified School District, it is crucial to explicitly address and demonstrate compliance with each of the listed RFP requirements, integrating them naturally into the relevant sections of the proposal."
                    )
            except FileNotFoundError:
                print("Warning: natomas_school_district_rfp_requirements.txt not found. Proceeding without specific RFP context for Natomas.")

        model = genai.GenerativeModel('gemini-1.5-flash')
        company = "Music Science & Technology Group" # Updated here
        client = f"{district} School District"
        project = "Educational Learning Outreach Program (ELOP)"
        services = "Music Integration, S.T.E.A.M. Education, Wellness Programs, Student Engagement Activities"
        benefits = "Enhanced Student Engagement, Improved Academic Performance, Increased Wellness, Community Involvement"
        cta = "We look forward to the opportunity to collaborate and make a meaningful impact together."
        today = date.today()
        year = today.year
        
        school_locations = ", ".join(selected_schools) if selected_schools else "Selected school sites"

        # --- Format cost_proposal here before inserting into the prompt template ---
        try:
            numeric_cost_proposal = float(cost_proposal)
            formatted_cost_proposal = f"${numeric_cost_proposal:,.2f}"
        except (ValueError, TypeError):
            formatted_cost_proposal = "$0.00" # Default or handle error appropriately
            print(f"Warning: Invalid cost_proposal value received: {cost_proposal}. Defaulting to $0.00.")


        # --- Read prompt template from file ---
        with open(os.path.join(app.config['INPUT_FILES_FOLDER'], 'proposal_prompt.txt'), 'r') as f:
            prompt_template = f.read()

        # Format the prompt using a dictionary for clarity with multiple variables
        prompt = prompt_template.format(
            district=district,
            today=today,
            days_per_week=days_per_week,
            num_weeks=num_weeks,
            school_locations=school_locations,
            formatted_cost_proposal=formatted_cost_proposal, # Pass the pre-formatted cost here
            year=year,
            about_msg=about_msg,
            proposal_context=proposal_context,
            all_districts_info=all_districts_info, # Although commented out in prompt, good practice to pass if template changes
            natomas_rfp_requirements_context_formatted=natomas_rfp_requirements_context,
            natomas_rfp_instruction_formatted=natomas_rfp_instruction_formatted,
            client=client,
            project=project,
            services=services,
            benefits=benefits,
            cta=cta
        )

        print("Generating proposal using Gemini model...")
        response = model.generate_content(prompt)
        text = normalize_empty_lines(response.text)
        
        document = Document()
        create_document_header(document)
        converter = MarkdownToDocxConverter(document=document)
        converter.convert(text)

        safe_district_name = re.sub(r'[^a-zA-Z0-9_]', '', district).replace(" ", "_")
        timestamp = int(time.time())
        filename = f"Proposal_{safe_district_name}_{timestamp}.docx"
        
        # Save to the DOWNLOAD_FOLDER
        output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
        document.save(output_path)
        print(f"Successfully created '{filename}' in '{app.config['DOWNLOAD_FOLDER']}' with custom header.")

        # Manage file rotation for this district
        manage_file_rotation(district, app.config['DOWNLOAD_FOLDER'])

        return text, filename
    except Exception as e:
        error_text = f"An error occurred while generating the proposal: {e}"
        print(error_text)
        print("##################################")
        print(INPUT_FILES_FOLDER)
        print(DOWNLOAD_FOLDER)
        print("##################################")

        return error_text, None

def get_school_data():
    # Assuming data.csv is in the root directory for now
    df = pd.read_csv("data.csv") 
    df.columns = df.columns.str.strip()
    return df

@app.route('/')
def index():
    df = get_school_data()
    districts = df['District Name'].unique().tolist()
    all_data_json = df.to_json(orient='records')
    return render_template('index.html', districts=districts, all_data=all_data_json)

@app.route('/proposal', methods=['POST'])
def generate_proposal():
    district_name = request.form.get('district')
    cost_proposal = request.form.get('cost_proposal')
    num_weeks = request.form.get('num_weeks')
    days_per_week = request.form.get('days_per_week')
    selected_schools = request.form.getlist('schoolname')

    proposal_text, filename = generate_proposal_from_row(
        district=district_name, 
        cost_proposal=cost_proposal, 
        num_weeks=num_weeks, 
        days_per_week=days_per_week, 
        selected_schools=selected_schools
    )

    proposal_data = {
        'district': district_name,
        'cost_proposal': cost_proposal,
        'school_name': ", ".join(selected_schools),
        'proposal_text': proposal_text,
        'filename': filename
    }
    return render_template('proposal.html', data=proposal_data)

@app.route('/download/<path:filename>')
def download_file(filename):
    safe_filename = secure_filename(filename)
    # Serve from the DOWNLOAD_FOLDER
    download_dir = app.config['DOWNLOAD_FOLDER']
    file_path = os.path.join(download_dir, safe_filename)
    
    if not os.path.normpath(file_path).startswith(download_dir) or not os.path.isfile(file_path):
        abort(404)

    return send_from_directory(download_dir, safe_filename, as_attachment=True)

@app.route('/stopServer', methods=['GET'])
def stopServer():
    os.kill(os.getpid(), signal.SIGINT)
    return jsonify({ "success": True, "message": "Server is shutting down..." })

@app.route("/healthz")
def healthz():
    return "ok", 200

# Initialize the application
_initialize()

if __name__ == '__main__':
    # Initialize the application
    #_initialize()

    """
    if env.get("DEPL") == "PROD":
        app.run(host='0.0.0.0', port=443, debug=True, ssl_context=(
            '/etc/letsencrypt/live/qaproposalmusicsciencegroup.com/fullchain.pem',
            '/etc/letsencrypt/live/qaproposalmusicsciencegroup.com/privkey.pem'
        ))
    else:
        app.run(host='0.0.0.0', port=5000, debug=True)
    """
    app.run(host='0.0.0.0', port=5000, debug=True)
