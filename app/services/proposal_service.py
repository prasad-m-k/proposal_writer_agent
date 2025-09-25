"""
Proposal Generation Service Module
"""
import os
import re
import time
import pandas as pd
from datetime import date, datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
import yaml
from jinja2 import Template
from docx2pdf import convert

from convert import MarkdownToDocxConverter
from config.settings import RFP_TYPE_FILES
import word_formatter


class GeminiAPIManager:
    """Manages Gemini API keys with smart switching and cooldown logic"""

    def __init__(self, primary_key, alternate_key):
        self.primary_key = primary_key
        self.alternate_key = alternate_key
        self.current_key = primary_key
        self.is_using_alternate = False
        self.last_switch_time = None
        self.switch_cooldown_hours = 1

    def can_switch_keys(self):
        """Check if we can switch keys (not switched in last hour)"""
        if self.last_switch_time is None:
            return True

        time_since_switch = datetime.now() - self.last_switch_time
        return time_since_switch >= timedelta(hours=self.switch_cooldown_hours)

    def switch_to_alternate_key(self):
        """Switch to alternate API key if cooldown period has passed"""
        if not self.can_switch_keys():
            print(f"Cannot switch keys - last switch was less than {self.switch_cooldown_hours} hour(s) ago")
            return False

        if not self.is_using_alternate:
            self.current_key = self.alternate_key
            self.is_using_alternate = True
            self.last_switch_time = datetime.now()
            genai.configure(api_key=self.current_key)
            print("Switched to alternate Gemini API key")
            return True
        else:
            print("Already using alternate key")
            return False

    def get_current_key_info(self):
        """Get info about current key usage"""
        key_type = "ALTERNATE" if self.is_using_alternate else "PRIMARY"
        cooldown_status = "Available" if self.can_switch_keys() else f"Cooldown until {self.last_switch_time + timedelta(hours=self.switch_cooldown_hours)}"
        return f"Using {key_type} key. Switch {cooldown_status}"


class ProposalService:
    """Service class for handling proposal generation"""
    
    def __init__(self, app_config):
        self.config = app_config
        self._configure_genai()
    
    def _configure_genai(self):
        """Configure Google Generative AI with fallback API key management"""
        primary_key = self.config.get('GEMINI_API_KEY')
        alternate_key = self.config.get('GEMINI_API_KEY_ALTERNATE')

        if primary_key and alternate_key:
            self.api_manager = GeminiAPIManager(primary_key, alternate_key)
            genai.configure(api_key=primary_key)
            print("Gemini API configured with primary key and fallback enabled")
        else:
            print("Warning: Both GEMINI_API_KEY and GEMINI_API_KEY_ALTERNATE are required for fallback functionality")
            if primary_key:
                genai.configure(api_key=primary_key)
                self.api_manager = None

    def call_gemini_with_fallback(self, prompt):
        """Call Gemini API with automatic fallback to alternate key on rate limit errors"""
        model = genai.GenerativeModel('gemini-2.0-flash')

        try:
            print(f"Attempting Gemini API call with {'alternate' if hasattr(self, 'api_manager') and self.api_manager.is_using_alternate else 'primary'} key...")
            response = model.generate_content(prompt)
            return response.text

        except Exception as e:
            error_message = str(e).lower()
            print(f"Gemini API error: {e}")

            # Check if it's a rate limit or quota error and we can try fallback
            is_rate_limit_error = any(keyword in error_message for keyword in [
                'rate limit', 'quota', 'resource exhausted', 'too many requests', 'limit exceeded'
            ])

            if is_rate_limit_error and hasattr(self, 'api_manager') and self.api_manager:
                print("Detected rate limit error. Attempting to switch to alternate API key...")

                if self.api_manager.switch_to_alternate_key():
                    try:
                        print("Retrying with alternate API key...")
                        model = genai.GenerativeModel('gemini-2.0-flash')
                        response = model.generate_content(prompt)
                        return response.text

                    except Exception as fallback_error:
                        print(f"Alternate API key also failed: {fallback_error}")
                        raise fallback_error

                else:
                    print("Could not switch to alternate key (cooldown active or already using alternate)")
                    raise e

            else:
                # Re-raise the original exception if it's not a rate limit error or no fallback available
                raise e

    def get_api_status(self):
        """Get current API key status information"""
        if hasattr(self, 'api_manager') and self.api_manager:
            return self.api_manager.get_current_key_info()
        else:
            return "Using single API key (no fallback configured)"

    def load_rfp_files(self, rfp_type):
        """Load the appropriate prompt and requirements files based on RFP type."""
        if rfp_type not in RFP_TYPE_FILES:
            rfp_type = "Extended Learning Opportunities Program"  # Default fallback

        files = RFP_TYPE_FILES[rfp_type]

        # Check if this RFP type uses YAML + Jinja strategy
        if 'yaml' in files and 'jinja' in files:
            return self.load_yaml_jinja_template(files)

        try:
            # Load requirements file
            requirements_path = os.path.join(self.config['INPUT_FILES_FOLDER'], files["requirements"])
            with open(requirements_path, 'r', encoding='utf-8') as f:
                print(f"reading from requirements_path: {requirements_path} ")
                requirements_context = f.read()

            # Load prompt file
            prompt_path = os.path.join(self.config['INPUT_FILES_FOLDER'], files["prompt"])
            with open(prompt_path, 'r', encoding='utf-8') as f:
                print(f"reading from prompt_path: {prompt_path} ")
                prompt_template = f.read()

            return prompt_template, requirements_context

        except FileNotFoundError as e:
            print(f"Warning: Could not load RFP files for {rfp_type}: {e}")
            # Fallback to default files
            try:
                with open(os.path.join(self.config['INPUT_FILES_FOLDER'], 'proposal_prompt.txt'), 'r', encoding='utf-8') as f:
                    prompt_template = f.read()
                with open(os.path.join(self.config['INPUT_FILES_FOLDER'], 'natomas_school_district_rfp_requirements.txt'), 'r', encoding='utf-8') as f:
                    requirements_context = f.read()
                return prompt_template, requirements_context
            except FileNotFoundError:
                raise ValueError(f"Could not load default prompt files for RFP type: {rfp_type}")

    def load_yaml_jinja_template(self, files):
        """Load YAML configuration and Jinja template for structured RFP generation"""
        try:
            # Load YAML configuration
            yaml_path = os.path.join(self.config['INPUT_FILES_FOLDER'], files["yaml"])
            with open(yaml_path, 'r', encoding='utf-8') as f:
                print(f"reading from yaml_path: {yaml_path}")
                yaml_config = yaml.safe_load(f)

            # Load Jinja template
            jinja_path = os.path.join(self.config['INPUT_FILES_FOLDER'], files["jinja"])
            with open(jinja_path, 'r', encoding='utf-8') as f:
                print(f"reading from jinja_path: {jinja_path}")
                jinja_template = f.read()

            # Return a dictionary that indicates YAML+Jinja mode
            return {
                'mode': 'yaml_jinja',
                'yaml_config': yaml_config,
                'jinja_template': jinja_template
            }

        except FileNotFoundError as e:
            print(f"Warning: Could not load YAML/Jinja files: {e}")
            raise ValueError(f"Could not load YAML/Jinja template files: {e}")
    
    def load_context_files(self):
        """Load common context files"""
        context_files = {}
        
        try:
            with open(os.path.join(self.config['INPUT_FILES_FOLDER'], 'district.csv'), 'r', encoding='utf-8') as f:
                context_files['all_districts_info'] = f.read()

            with open(os.path.join(self.config['INPUT_FILES_FOLDER'], 'about_mstg.txt'), 'r', encoding='utf-8') as f:
                context_files['about_mstg'] = f.read()

            with open(os.path.join(self.config['INPUT_FILES_FOLDER'], 'about_minkh.txt'), 'r', encoding='utf-8') as f:
                context_files['about_minkh'] = f.read()
                
        except FileNotFoundError as e:
            print(f"Warning: Could not load context file: {e}")
            
        return context_files
    
    def create_document_header(self, document):
        """Add a pre-defined header to the document"""
        LOGO_PATH = 'static/assets/mstg_large_logo.png'
        ADDRESS_LINES = [
            ("Musical Instruments N Kids Hands", True, 8),
            ("Music Science & Technology Group", True, 8),
            ("2324 L St, STE 309, Sacramento, CA 95816", False, 8),
            ("Ph. (216) 903-3756", False, 8)
        ]

        section = document.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = ""

        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(5.0)

        logo_cell = header_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        
        if os.path.exists(LOGO_PATH):
            logo_run.add_picture(LOGO_PATH, height=Inches(0.75))

        address_cell = header_table.cell(0, 1)
        address_paragraph = address_cell.paragraphs[0]
        address_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for text, is_bold, font_size in ADDRESS_LINES:
            run = address_paragraph.add_run(text + '\n')
            run.bold = is_bold
            font = run.font
            font.size = Pt(font_size)
    
    def normalize_empty_lines(self, text):
        """Reduce any sequence of multiple newlines to single newlines"""
        return re.sub(r'\n{2,}', r'\n', text)

    def fill_form_fields(self, text, prompt_variables):
        """Fill in form fields with actual company information and add signature lines"""
        # Replace underscored form fields with actual data
        replacements = {
            r'\*\*Company Name:\*\* _{10,}': f"**Company Name:** {prompt_variables.get('company_name', 'Musical Instruments N Kids Hands (M.I.N.K.H.) - Music Science & Technology Group (MSTG)')}",
            r'\*\*Authorized Representative Name:\*\* _{10,}': f"**Authorized Representative Name:** {prompt_variables.get('representative_name', 'A.P. Moore, Program Coordinator')}",
            r'\*\*Title:\*\* _{10,}': f"**Title:** {prompt_variables.get('representative_title', 'Program Coordinator')}",
            r'\*\*Email:\*\* _{10,}': f"**Email:** {prompt_variables.get('company_email', 'aphilanda@musicsciencegroup.com')}",
            r'\*\*Phone:\*\* _{10,}': f"**Phone:** {prompt_variables.get('company_phone', '(216) 903-3756')}",
            r'\*\*Address:\*\* _{10,}': f"**Address:** {prompt_variables.get('company_address', '2150 Capitol Avenue Sacramento, CA 95816')}",
            r'\*\*Authorized Signature:\*\* _{10,}': f"**Authorized Signature:** {prompt_variables.get('representative_name', 'A.P. Moore, Program Coordinator')}",
            r'\*\*Date:\*\* _{10,}': f"**Date:** {prompt_variables.get('today', date.today())}"
        }

        for pattern, replacement in replacements.items():
            text = re.sub(pattern, replacement, text)

        # Add signature lines for form fields that might have been filled by AI
        signature_patterns = {
            # Pattern: "I, Name, Title" -> "I, Name _________________________, Title _______________________"
            r'I, ([^,]+), ([^\n]+)\n\(Name\)\s+\(Title\)': r'I, \1 _________________________, \2 _______________________  \n(Name)                                   (Title)',

            # Pattern: "Of Company Name hereby certify:" -> "Of Company Name ________________________________ hereby certify:"
            r'Of ([^\n]+) hereby certify:\n\(Company Name\)': r'Of \1 ______________________________________ hereby certify:  \n(Company Name)',

            # Pattern: "**Signature:** Name" -> "**Signature:** ____________________________"
            r'\*\*Signature:\*\* ([^\n]+)': r'**Signature:** ____________________________',

            # Pattern: "**Authorized Signature:** Name" -> "**Authorized Signature (Disclosee):** ____________________________"
            r'\*\*Authorized Signature \(Disclosee\):\*\* ([^\n]+)': r'**Authorized Signature (Disclosee):** ____________________________',

            # Pattern: "**Date:** date" -> "**Date:** _____________"
            r'\*\*Date:\*\* [0-9]{4}-[0-9]{2}-[0-9]{2}': r'**Date:** _____________',
            r'\*\*Date:\*\* [^\n]+\d{4}': r'**Date:** _____________',

            # Pattern: "**Initial Here:** letters" -> "**Initial Here:** ____"
            r'\*\*Initial Here:\*\* ([A-Z]+)': r'**Initial Here:** ____',
        }

        # Apply signature line patterns
        for pattern, replacement in signature_patterns.items():
            text = re.sub(pattern, replacement, text)

        return text
    
    def clean_and_format_currency(self, value, default="$0.00"):
        """Clean and format currency values"""
        try:
            # Handle empty or None values
            if not value or str(value).strip() == '':
                return default

            clean_value = re.sub(r'[^\d.-]', '', str(value))  # Allow negative sign

            if not clean_value or clean_value == '.':
                return default

            numeric_value = float(clean_value)

            # Use locale-independent formatting
            formatted_result = "${:,.2f}".format(numeric_value)
            return formatted_result
        except (ValueError, TypeError) as e:
            print(f"Warning: Invalid currency value received: {value}. Error: {e}. Defaulting to {default}.")
            return default
    
    def prepare_prompt_variables(self, **kwargs):
        """Prepare variables for the prompt template"""
        # Load context files
        context_files = self.load_context_files()
        
        # Clean and format financial values
        formatted_cost_proposal = self.clean_and_format_currency(kwargs.get('cost_proposal'))

        # Calculate actual cost per student based on fixed program structure
        # Program: 439 students, 7 weeks, $75,000 total cost
        ##program_total_cost = 75000.00
        ##total_students = 439
        ##actual_cost_per_student = program_total_cost / total_students if total_students > 0 else 0

        # Use calculated value or form value, prefer calculated
        cost_per_student_value = kwargs.get('cost_per_student')
        if cost_per_student_value and float(str(cost_per_student_value).replace('$', '').replace(',', '')) > 0:
            formatted_cost_per_student = self.clean_and_format_currency(cost_per_student_value)
        else:
            formatted_cost_per_student = self.clean_and_format_currency("0")

        formatted_daily_cost = self.clean_and_format_currency(kwargs.get('daily_cost'))
        formatted_weekly_cost = self.clean_and_format_currency(kwargs.get('weekly_cost'))

        # Debug output
        print(f"DEBUG: cost_per_student from form: {kwargs.get('cost_per_student')}")
        print(f"DEBUG: formatted_cost_per_student: {formatted_cost_per_student}")
        print(f"DEBUG: daily_cost from form: {kwargs.get('daily_cost')}")
        print(f"DEBUG: weekly_cost from form: {kwargs.get('weekly_cost')}")
        print(f"DEBUG: cost_per_school from form: {kwargs.get('cost_per_school')}")
        print(f"DEBUG: cost_proposal from form: {kwargs.get('cost_proposal')}")
        
        # Prepare school locations
        selected_schools = kwargs.get('selected_schools', [])
        school_locations = ", ".join(selected_schools) if selected_schools else "Selected school sites"

        # Parse selected schools list if available
        selected_schools_list = kwargs.get('selected_schools_list')
        if selected_schools_list:
            try:
                import json
                parsed_schools = json.loads(selected_schools_list)
                school_locations = ", ".join(parsed_schools) if parsed_schools else school_locations
            except (json.JSONDecodeError, TypeError):
                # Fallback to the original list
                pass

        # Format cost per school
        formatted_cost_per_school = self.clean_and_format_currency(kwargs.get('cost_per_school'))

        # Format program start date
        program_start_date = kwargs.get('program_start_date')
        if program_start_date:
            try:
                from datetime import datetime
                date_obj = datetime.strptime(program_start_date, '%Y-%m-%d')
                formatted_start_date = date_obj.strftime('%B %d, %Y')  # e.g., "October 07, 2024"

                # Calculate end date (assuming 7 weeks)
                from datetime import timedelta
                end_date_obj = date_obj + timedelta(weeks=7)
                formatted_end_date = end_date_obj.strftime('%B %d, %Y')
                program_dates = f"{formatted_start_date} - {formatted_end_date}"
            except (ValueError, TypeError):
                program_dates = "January 13, 2025 - March 07, 2025"
        else:
            program_dates = "January 13, 2025 - March 07, 2025"
        
        # Format RFP instructions
        district = kwargs.get('district', 'N/A')
        rfp_type = kwargs.get('rfp_type', 'Extended Learning Opportunities Program')
        
        if rfp_type != "Extended Learning Opportunities Program":
            rfp_instruction_formatted = (
                f"\nWhen generating the proposal for {district}, it is crucial to explicitly address and demonstrate compliance with each of the listed RFP requirements for {rfp_type}, integrating them naturally into the relevant sections of the proposal."
            )
        else:
            rfp_instruction_formatted = ""
            if district == "Natomas Unified School District":
                rfp_instruction_formatted = (
                    "\nWhen generating the proposal for Natomas Unified School District, it is crucial to explicitly address and demonstrate compliance with each of the listed RFP requirements, integrating them naturally into the relevant sections of the proposal."
                )
        
        # Company information variables
        company_name = kwargs.get('company_name', 'Musical Instruments N Kids Hands (M.I.N.K.H.) - Music Science & Technology Group (MSTG)')
        representative_name = kwargs.get('representative_name', 'A.P. Moore, Program Coordinator')
        representative_title = kwargs.get('representative_title', 'Program Coordinator')
        company_email = kwargs.get('company_email', 'aphilanda@musicsciencegroup.com')
        company_phone = kwargs.get('company_phone', '(216) 903-3756')
        company_address = kwargs.get('company_address', '2150 Capitol Avenue Sacramento, CA 95816')

        # Create comprehensive variables dictionary
        prompt_variables = {
            'district': district,
            'rfp_type': rfp_type,
            'today': date.today(),
            'days_per_week': kwargs.get('days_per_week', 'N/A'),
            'num_weeks': kwargs.get('num_weeks', 'N/A'),
            'hours_per_day': kwargs.get('hours_per_day', '3'),
            'school_locations': school_locations,
            'formatted_cost_proposal': formatted_cost_proposal,
            'formatted_daily_cost': formatted_daily_cost,
            'formatted_weekly_cost': formatted_weekly_cost,
            'total_students': kwargs.get('total_students', 'N/A'),
            'formatted_cost_per_student': formatted_cost_per_student,
            'formatted_cost_per_school': formatted_cost_per_school,
            'program_dates': program_dates,
            'year': date.today().year,
            'rfp_instruction_formatted': rfp_instruction_formatted,
            'client': f"{district} School District",
            'project': {rfp_type},
            'services': "Music Integration, S.T.E.A.M. Education, Wellness Programs, Student Engagement Activities",
            'benefits': "Enhanced Student Engagement, Improved Academic Performance, Increased Wellness, Community Involvement",
            'cta': "We look forward to the opportunity to collaborate and make a meaningful impact together.",

            # Company information
            'company_name': company_name,
            'representative_name': representative_name,
            'representative_title': representative_title,
            'company_email': company_email,
            'company_phone': company_phone,
            'company_address': company_address
        }
        
        # Add context files
        prompt_variables.update(context_files)
        
        # Add RFP requirements context
        rfp_data = self.load_rfp_files(rfp_type)
        if isinstance(rfp_data, dict) and rfp_data.get('mode') == 'yaml_jinja':
            # For YAML+Jinja mode, use the RFP type as context
            requirements_context = f"RFP Type: {rfp_type} (using YAML+Jinja template)"
        else:
            # Traditional mode - extract requirements context
            _, requirements_context = rfp_data
        prompt_variables['rfp_requirements_context_formatted'] = requirements_context
        
        # For backward compatibility
        prompt_variables.update({
            'natomas_rfp_requirements_context_formatted': requirements_context,
            'natomas_rfp_instruction_formatted': rfp_instruction_formatted
        })

        return prompt_variables
    
    def generate_proposal_text(self, template_data, prompt_variables):
        """Generate proposal text using AI - handles both traditional and YAML+Jinja approaches"""
        try:
            # Check if template_data is a YAML+Jinja config dictionary
            if isinstance(template_data, dict) and template_data.get('mode') == 'yaml_jinja':
                # YAML + Jinja approach
                return self.generate_with_yaml_jinja(template_data, prompt_variables)
            else:
                # Traditional approach - extract prompt template from tuple
                if isinstance(template_data, tuple):
                    prompt_template, _ = template_data
                else:
                    prompt_template = template_data  # Fallback for direct string

                prompt = prompt_template.format(**prompt_variables)

                print(f"Generating proposal using Gemini model for RFP type: {prompt_variables.get('rfp_type')}...")
                ai_text_raw = self.call_gemini_with_fallback(prompt)

                # Post-process the AI text to fill in form fields with actual company data
                ai_text = self.normalize_empty_lines(ai_text_raw)
                ai_text = self.fill_form_fields(ai_text, prompt_variables)

                return ai_text

        except Exception as e:
            error_text = f"An error occurred while generating the proposal text: {e}"
            print(error_text)
            return error_text

    def generate_with_yaml_jinja(self, config_data, prompt_variables):
        """Generate proposal using YAML configuration and Jinja template"""
        try:
            # Extract YAML config and Jinja template from the config_data
            yaml_config = config_data['yaml_config']
            jinja_template_content = config_data['jinja_template']

            # Populate YAML variables with form data
            populated_vars = self.populate_yaml_vars(yaml_config['vars'], prompt_variables)

            # Render the Jinja template with populated variables
            template = Template(jinja_template_content)
            rendered_content = template.render(**populated_vars)

            # Create enhanced prompt for Gemini
            enhanced_prompt = self.create_enhanced_prompt_for_yaml(rendered_content, prompt_variables)

            rfp_type = prompt_variables.get('rfp_type', 'Request for Qualifications')
            print(f"Generating proposal using YAML+Jinja strategy for RFP type: {rfp_type}...")
            ai_text_raw = self.call_gemini_with_fallback(enhanced_prompt)

            return self.normalize_empty_lines(ai_text_raw)

        except Exception as e:
            error_text = f"An error occurred in YAML+Jinja generation: {e}"
            print(error_text)
            return error_text

    def populate_yaml_vars(self, yaml_vars, prompt_variables):
        """Populate YAML variables with form data"""
        # Deep copy the YAML structure to avoid modifying the original
        import copy
        populated_vars = copy.deepcopy(yaml_vars)

        # Populate submission information
        if 'submission' in populated_vars:
            populated_vars['submission']['company_name'] = prompt_variables.get('company_name', '')
            populated_vars['submission']['rep_name'] = prompt_variables.get('representative_name', '')
            populated_vars['submission']['title'] = prompt_variables.get('representative_title', '')
            populated_vars['submission']['email'] = prompt_variables.get('company_email', '')
            populated_vars['submission']['phone'] = prompt_variables.get('company_phone', '')
            populated_vars['submission']['address'] = prompt_variables.get('company_address', '')
            populated_vars['submission']['signature_name'] = prompt_variables.get('representative_name', '')
            populated_vars['submission']['signature_date'] = str(prompt_variables.get('today', ''))

        # Populate budget information if available
        if 'budget' in populated_vars and 'categories' in populated_vars['budget']:
            # Extract costs from form data if available (you may need to enhance this based on your form structure)
            total_cost = prompt_variables.get('cost_proposal', '$0.00')
            if isinstance(total_cost, str):
                try:
                    # Clean currency and convert to float
                    clean_cost = float(re.sub(r'[^\d.-]', '', total_cost))
                    # Distribute cost across categories (basic allocation - you can enhance this)
                    populated_vars['budget']['categories']['staffing'] = clean_cost * 0.6  # 60% staffing
                    populated_vars['budget']['categories']['instructional_materials'] = clean_cost * 0.1
                    populated_vars['budget']['categories']['program_supplies'] = clean_cost * 0.1
                    populated_vars['budget']['categories']['supervision'] = clean_cost * 0.1
                    populated_vars['budget']['categories']['professional_development'] = clean_cost * 0.05
                    populated_vars['budget']['categories']['transportation'] = clean_cost * 0.025
                    populated_vars['budget']['categories']['admin_costs'] = clean_cost * 0.025
                except (ValueError, TypeError):
                    pass  # Keep default values if conversion fails

        # Populate legal information
        if 'legal' in populated_vars:
            populated_vars['legal']['nda_disclosee_name'] = prompt_variables.get('company_name', '')
            populated_vars['legal']['workers_comp_company_name'] = prompt_variables.get('company_name', '')
            populated_vars['legal']['workers_comp_rep_name'] = prompt_variables.get('representative_name', '')

        # Set up references with default data (can be enhanced to pull from form)
        if 'references' in populated_vars:
            default_refs = [
                {
                    'organization': 'Natomas Unified School District',
                    'contact_name': 'Program Administrator',
                    'phone': '(916) 567-5000',
                    'email': 'programs@natomas.net',
                    'dates_of_service': '2023-2024',
                    'description': 'Extended Learning Opportunities Program'
                },
                {
                    'organization': 'Sacramento City Unified School District',
                    'contact_name': 'Enrichment Coordinator',
                    'phone': '(916) 643-7400',
                    'email': 'enrichment@scusd.edu',
                    'dates_of_service': '2022-2023',
                    'description': 'Music & STEAM Integration Program'
                },
                {
                    'organization': 'Elk Grove Unified School District',
                    'contact_name': 'After School Programs Manager',
                    'phone': '(916) 686-7700',
                    'email': 'afterschool@egusd.net',
                    'dates_of_service': '2021-2022',
                    'description': 'Arts & Wellness Programs'
                }
            ]

            for i in range(min(3, len(populated_vars['references']))):  # Limit to first 3 references
                if i < len(default_refs):
                    populated_vars['references'][i].update(default_refs[i])

        return populated_vars

    def create_enhanced_prompt_for_yaml(self, rendered_content, prompt_variables):
        """Create an enhanced prompt for Gemini using the rendered Jinja content"""

        # Load context information
        context_files = self.load_context_files()

        enhanced_prompt = f"""
You are a proposal writer completing an RFP response for {prompt_variables.get('district', 'N/A')} School District. This is a CRITICAL GOVERNMENT CONTRACT submission that must follow EXACT formatting requirements.

WARNING: School districts have STRICT documentation order requirements. ANY deviation from the provided template structure will result in AUTOMATIC DISQUALIFICATION.

ORGANIZATION CONTEXT:
{context_files.get('about_mstg', 'Musical Instruments N Kids Hands (M.I.N.K.H.) - Music Science & Technology Group (MSTG) specializes in music integration and STEAM education programs.')}

{context_files.get('about_minkh', 'M.I.N.K.H. focuses on hands-on learning experiences that combine music, science, technology, engineering, arts, and mathematics.')}

PROJECT DETAILS:
- Client: {prompt_variables.get('district', 'N/A')} School District
- Program: {prompt_variables.get('rfp_type', 'Extended Learning Opportunities Program')}
- Students: {prompt_variables.get('total_students', 'N/A')}
- Duration: {prompt_variables.get('program_dates', 'N/A')}
- Investment: {prompt_variables.get('formatted_cost_proposal', 'N/A')}
- Cost/Student: {prompt_variables.get('formatted_cost_per_student', 'N/A')}

TEMPLATE TO COMPLETE (MAINTAIN EXACT FORMAT):
{rendered_content}

CRITICAL FORMATTING REQUIREMENTS:
1. DO NOT change any headings, section numbers, or appendix labels
2. DO NOT reorder any sections - keep the EXACT sequence provided
3. DO NOT add extra sections or remove any existing sections
4. DO NOT modify table structures, column headers, or formatting
5. DO NOT change markdown formatting (**, ##, |, ---, etc.)
6. DO NOT add introductory text, conclusions, or explanations outside the template

CONTENT REQUIREMENTS:
7. Replace ONLY "TBD" placeholders with specific, professional content
8. For narrative sections: Write 2-3 focused paragraphs demonstrating expertise
9. For application questions: Provide direct, substantive answers (3-5 sentences each)
10. Use measurable outcomes and specific examples from music/STEAM education
11. Maintain professional, confident tone throughout
12. Address diversity, equity, inclusion naturally within existing sections
13. Keep all budget figures, dates, and contact information as provided

COMPLIANCE CHECK:
- Every "TBD" must be replaced with appropriate content
- All headings must remain exactly as shown
- All tables must maintain their structure
- All legal text sections must remain unchanged
- Document must be ready for direct conversion to PDF/DOCX

HALLUCINATION PREVENTION:
- DO NOT invent new sections, appendices, or requirements not in the template
- DO NOT add creative flourishes or additional formatting beyond what's provided
- DO NOT include meta-commentary about the proposal process
- DO NOT mention AI assistance or generation in the content
- DO NOT add disclaimers, notes, or explanations anywhere in the document

RETURN REQUIREMENT:
Output ONLY the completed proposal template with all "TBD" placeholders filled. Start immediately with the first line of the template (the title). Do not include any explanations, notes, or additional text outside the template structure. The output must be ready for immediate submission to the school district.

FINAL CHECK BEFORE SUBMISSION:
1. Does the output start with the exact title from the template?
2. Are all headings preserved exactly as shown?
3. Are all table structures intact?
4. Are all "TBD" entries replaced with appropriate content?
5. Does the output end with the last section from the template?
"""

        return enhanced_prompt
    
    def get_document_title(self, district, rfp_type):
        """Generate appropriate document title based on RFP type and district"""
        title_mapping = {
            "Extended Learning Opportunities Program": f"ELOP Program Proposal for {district}",
            "Request for Qualifications": f"Request for Qualifications (RFQ)\nEnrichment Program Providers at {district}",
            "Summer School Core Program Providers": f"ELOP Program Proposal for Summer School Core Program Providers at {district}",
            "After School Core Program Providers": f"ELOP Program Proposal for After School Core Program Providers at {district}"
        }
        return title_mapping.get(rfp_type, f"Proposal for {district}")

    def create_document(self, text, district, rfp_type):
        """Create and save the Word document"""
        try:
            document = Document()
            self.create_document_header(document)

            # Add document title
            title = self.get_document_title(district, rfp_type)
            title_paragraph = document.add_heading(title, level=0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Convert text to document
            converter = MarkdownToDocxConverter(document=document)
            text = text.replace('** ', '**').replace(' **', '**')
            converter.convert(text)

            # Generate filename
            safe_district_name = re.sub(r'[^a-zA-Z0-9_]', '', district).replace(" ", "_")
            safe_rfp_type = re.sub(r'[^a-zA-Z0-9_]', '', rfp_type).replace(" ", "_")
            timestamp = int(time.time())
            filename = f"Proposal_{safe_district_name}_{safe_rfp_type}_{timestamp}.docx"
            
            output_path = os.path.join(self.config['DOWNLOAD_FOLDER'], filename)
            document.save(output_path)

            # Format the document
            word_formatter.format_word_document(output_path, output_path)

            # Generate PDF from the DOCX
            pdf_filename = filename.replace('.docx', '.pdf')
            pdf_output_path = os.path.join(self.config['DOWNLOAD_FOLDER'], pdf_filename)

            try:
                convert(output_path, pdf_output_path)
                print(f"Successfully created PDF '{pdf_filename}' from DOCX")
            except Exception as pdf_error:
                print(f"Warning: Failed to create PDF: {pdf_error}")
                # Continue even if PDF creation fails

            print(f"Successfully created '{filename}' in '{self.config['DOWNLOAD_FOLDER']}' with custom header.")

            return filename
            
        except Exception as e:
            print(f"Error creating document: {e}")
            return None
    
    def manage_file_rotation(self, district_name, max_files=5):
        """Manage file rotation to keep only recent files"""
        safe_district_name = re.sub(r'[^a-zA-Z0-9_]', '', district_name).replace(" ", "_")
        
        file_pattern = rf"Proposal_{re.escape(safe_district_name)}_.*_(\d+)\.docx"
        
        district_files = []
        download_folder = self.config['DOWNLOAD_FOLDER']
        
        for filename in os.listdir(download_folder):
            match = re.match(file_pattern, filename)
            if match:
                timestamp = int(match.group(1))
                district_files.append((timestamp, os.path.join(download_folder, filename)))
                
        # Sort files by timestamp, oldest first
        district_files.sort(key=lambda x: x[0])

        # Delete old files if we have more than max_files
        if len(district_files) > max_files:
            files_to_delete = district_files[:len(district_files) - max_files]
            for _, filepath in files_to_delete:
                try:
                    os.remove(filepath)
                    print(f"Deleted old proposal file: {filepath}")
                except OSError as e:
                    print(f"Error deleting file {filepath}: {e}")
    
    def generate_proposal_text_only(self, **kwargs):
        """Generate only the proposal text without creating a document"""
        try:
            # Load appropriate files based on RFP type
            rfp_type = kwargs.get('rfp_type', 'Extended Learning Opportunities Program')
            rfp_data = self.load_rfp_files(rfp_type)

            # Prepare variables
            prompt_variables = self.prepare_prompt_variables(**kwargs)

            # Generate proposal text
            proposal_text = self.generate_proposal_text(rfp_data, prompt_variables)

            return proposal_text

        except Exception as e:
            error_text = f"An error occurred while generating the proposal text: {e}"
            print(error_text)
            return error_text

    def generate_proposal(self, **kwargs):
        """Main method to generate a complete proposal"""
        try:
            # Load appropriate files based on RFP type
            rfp_type = kwargs.get('rfp_type', 'Extended Learning Opportunities Program')
            rfp_data = self.load_rfp_files(rfp_type)

            # Prepare variables
            prompt_variables = self.prepare_prompt_variables(**kwargs)

            # Generate proposal text
            proposal_text = self.generate_proposal_text(rfp_data, prompt_variables)

            # Create document
            district = kwargs.get('district', 'N/A')
            filename = self.create_document(proposal_text, district, rfp_type)

            # Manage file rotation
            if filename:
                self.manage_file_rotation(district)

            return proposal_text, filename

        except Exception as e:
            error_text = f"An error occurred while generating the proposal: {e}"
            print(error_text)
            return error_text, None



class DataService:
    """Service class for handling data operations"""
    
    def __init__(self, app_config):
        self.config = app_config
    
    def get_school_data(self):
        """Read and return school data from CSV"""
        try:
            df = pd.read_csv(os.path.join(self.config['INPUT_FILES_FOLDER'], "data.csv"))
            df.columns = df.columns.str.strip()
            return df
        except Exception as e:
            print(f"Error loading school data: {e}")
            return pd.DataFrame()
    
    def get_districts(self):
        """Get list of unique districts"""
        df = self.get_school_data()
        return df['District Name'].unique().tolist() if not df.empty else []
