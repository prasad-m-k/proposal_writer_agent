# convert.py

import re
from docx import Document
from docx.shared import Inches

class MarkdownToDocxConverter:
    """
    A class to convert a Markdown formatted string into a .docx file.
    Can be initialized with an existing Document object to append content.
    """

    def __init__(self, document=None):
        """
        Initializes the converter.
        Args:
            document (Document, optional): An existing python-docx Document object. 
                                           If None, a new one is created.
        """
        if document:
            self.doc = document
        else:
            self.doc = Document()

    def convert(self, markdown_text):
        """
        Main conversion method. Adds markdown text to the document object.
        Note: This method no longer saves the file.
        
        Args:
            markdown_text (str): The string containing Markdown text.
        """
        self.inline_pattern = re.compile(
            r'(\*{1,3}|_{1,3}|~~|`)'
            r'(.+?)'
            r'\1'
        )
        lines = markdown_text.strip().split('\n')

        for line in lines:
            if line.startswith('#'):
                self._add_heading(line)
            elif line.strip() in ['---', '***', '___']:
                self._add_horizontal_rule()
            elif line.startswith('>'):
                self._add_blockquote(line)
            elif re.match(r'^\s*(\*|-|\+)\s', line) or re.match(r'^\s*\d+\.\s', line):
                self._add_list_item(line)
            elif line.strip() == '':
                self.doc.add_paragraph()
            else:
                self._add_formatted_paragraph(line)
        # The document is NOT saved here. The calling function is responsible for saving.

    # ... (the rest of the _add_... methods remain exactly the same) ...

    def _add_heading(self, line):
        """Adds a heading based on the number of '#' characters."""
        level = 0
        while line[level] == '#':
            level += 1
        
        heading_text = line[level:].strip()
        self.doc.add_heading(heading_text, level=min(level, 6))

    def _add_horizontal_rule(self):
        """Adds a horizontal rule to the document."""
        p = self.doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.element.get_or_add_pBdr().append(
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        )

    def _add_blockquote(self, line):
        """Adds a blockquote, removing the '>' marker."""
        quote_text = line.lstrip('> ').strip()
        self.doc.add_paragraph(quote_text, style='Quote')

    def _add_list_item(self, line):
        """Adds a list item, handling indentation and type (bullet/number)."""
        indent_level = len(line) - len(line.lstrip())
        indent_spaces = indent_level // 2
        style = 'List Bullet'
        if re.match(r'^\s*\d+\.\s', line):
            style = 'List Number'
        if indent_spaces > 0:
            style = f"{style} {indent_spaces + 1}"
        text = re.sub(r'^\s*(\*|-|\+)\s*|\s*\d+\.\s*', '', line).strip()
        self.doc.add_paragraph(text, style=style)

    def _add_formatted_paragraph(self, line):
        """Adds a paragraph with complex inline formatting."""
        p = self.doc.add_paragraph()
        last_end = 0
        for match in self.inline_pattern.finditer(line):
            start, end = match.span()
            marker = match.group(1)
            text = match.group(2)
            if start > last_end:
                p.add_run(line[last_end:start])
            run = p.add_run(text)
            if marker in ['**', '__', '***', '___']:
                run.bold = True
            if marker in ['*', '_', '***', '___']:
                run.italic = True
            if marker == '`':
                font = run.font
                font.name = 'Courier New'
            if marker == '~~':
                run.font.strike = True
            last_end = end
        if last_end < len(line):
            p.add_run(line[last_end:])